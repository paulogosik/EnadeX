# -*- coding: utf-8 -*-
"""
Gera o documento academico (.docx) do subprojeto ENADE-Time Distribuido,
seguindo a estrutura de 10 itens solicitada pelo orientador.

Todo o conteudo textual abaixo foi sintetizado a partir da documentacao real
do projeto (README.md, docs/ARQUITETURA.md, docs/DESIGN_LOG.md,
docs/DICIONARIO_VARIAVEIS.md, docs/INTEGRACAO_ECOSSISTEMA.md,
docs/LOG_EXECUCAO_ETL.md, docs/RESULTADOS_BENCHMARK.md,
docs/GUIA_APRESENTACAO.md, docs/EVIDENCIAS_TESTES.md, docs/PROXIMOS_PASSOS.md)
e do codigo-fonte efetivamente implementado — nao ha numeros nem trechos de
codigo inventados.

Números de benchmark: TODOS vêm do banco (views v_benchmark_resumo /
v_benchmark_metricas da campanha oficial) via dados_benchmark.carregar().
verificar_numeros.py confere o .docx gerado contra as views.

Uso:
    python docs/geradores/gerar_documento.py [--campanha <uuid>] [--ativos <pasta>]
    (--ativos: pasta com diagramas/ e screenshots/; padrão apresentacao/documento_final)
"""
from __future__ import annotations

import argparse
import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from dados_benchmark import br, carregar, data_extenso, pct, seg, sx  # noqa: E402

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent.parent
_ap = argparse.ArgumentParser()
_ap.add_argument("--campanha", default="oficial")
_ap.add_argument("--ativos", default=os.environ.get(
    "GERADORES_ATIVOS", str(RAIZ / "apresentacao" / "documento_final")))
_ap.add_argument("--saida", default=str(Path(__file__).resolve().parent / "out"))
ARGS = _ap.parse_args()
OUT_DIR = Path(ARGS.saida)
ATIVOS = Path(ARGS.ativos)
DIAGRAMAS = ATIVOS / "diagramas"
D = carregar(ARGS.campanha)

AZUL_ESCURO = RGBColor(0x1D, 0x2C, 0x46)
AZUL_MEDIO = RGBColor(0x2B, 0x44, 0x70)
CINZA_TEXTO = RGBColor(0x33, 0x33, 0x33)
CINZA_CLARO = RGBColor(0x66, 0x66, 0x66)

_QUADRO_N = [0]
_FIGURA_N = [0]
_TABELA_N = [0]


def prox(counter):
    counter[0] += 1
    return counter[0]


# ---------------------------------------------------------------------------
# Helpers de baixo nivel
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="999999", sz=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def h1(doc, texto):
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.page_break_before = True
    run = p.add_run(texto)
    run.font.color.rgb = AZUL_ESCURO
    return p


def h2(doc, texto):
    p = doc.add_heading(level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texto)
    run.font.color.rgb = AZUL_MEDIO
    return p


def h3(doc, texto):
    p = doc.add_heading(level=3)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL_MEDIO
    return p


def par(doc, texto="", *, bold=False, italic=False, size=11, color=None,
        space_after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if texto:
        r = p.add_run(texto)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color or CINZA_TEXTO
    return p


def rich_par(doc, partes, *, size=11, space_after=8):
    """partes: lista de (texto, dict_estilo). dict_estilo pode ter bold/italic."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for texto, estilo in partes:
        r = p.add_run(texto)
        r.font.size = Pt(size)
        r.font.color.rgb = CINZA_TEXTO
        r.bold = estilo.get("bold", False)
        r.italic = estilo.get("italic", False)
    return p


def bullets(doc, itens, *, size=11):
    for item in itens:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if isinstance(item, tuple):
            texto, negrito = item
            r = p.add_run(negrito)
            r.bold = True
            r.font.size = Pt(size)
            r2 = p.add_run(texto)
            r2.font.size = Pt(size)
        else:
            r = p.add_run(item)
            r.font.size = Pt(size)


def numbered(doc, itens, *, size=11):
    for item in itens:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if isinstance(item, tuple):
            titulo, resto = item
            r = p.add_run(titulo)
            r.bold = True
            r.font.size = Pt(size)
            r2 = p.add_run(resto)
            r2.font.size = Pt(size)
        else:
            r = p.add_run(item)
            r.font.size = Pt(size)


def legenda(doc, texto, *, space_before=4, space_after=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = CINZA_CLARO


def imagem(doc, caminho: Path, legenda_texto: str, *, largura_in=6.3):
    doc.add_picture(str(caminho), width=Inches(largura_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    n = prox(_FIGURA_N)
    legenda(doc, f"Figura {n} — {legenda_texto}")
    return n


def tabela(doc, headers, rows, widths_cm=None, *, legenda_texto=None, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htxt)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "2B4470")
        set_cell_borders(hdr[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(font_size)
            r.font.color.rgb = CINZA_TEXTO
            set_cell_borders(cells[i])
    if widths_cm:
        set_col_widths(t, widths_cm)
    if legenda_texto:
        n = prox(_TABELA_N)
        legenda(doc, f"Tabela {n} — {legenda_texto}", space_before=6)
    return t


def code_quadro(doc, titulo, codigo, explicacao, *, fonte=None):
    n = prox(_QUADRO_N)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Quadro {n} — {titulo}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = AZUL_MEDIO

    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, "1E1E2E")
    set_cell_borders(cell, color="0F0F1A", sz=6)
    cell.text = ""
    linhas = codigo.strip("\n").split("\n")
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    p0.paragraph_format.line_spacing = 1.0
    for i, ln in enumerate(linhas):
        target = p0 if i == 0 else cell.add_paragraph()
        target.paragraph_format.space_after = Pt(0)
        target.paragraph_format.line_spacing = 1.0
        run = target.add_run(ln if ln.strip() else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xDC)

    cap = f"Explicação: {explicacao}"
    if fonte:
        cap += f" (fonte: {fonte})"
    legenda(doc, cap, space_before=4, space_after=16)
    return n


def screenshot_real(doc, titulo, arquivo_png, comando_ou_url, *, largura_in=6.3):
    """Embute uma captura real (PNG em screenshots/) com legenda numerada."""
    caminho = ATIVOS / arquivo_png
    if not caminho.exists():
        par(doc, f"[captura ausente: {arquivo_png}]", italic=True, color=CINZA_CLARO)
        n = prox(_FIGURA_N)
        legenda(doc, f"Figura {n} — {titulo} ({comando_ou_url}).")
        return n
    doc.add_picture(str(caminho), width=Inches(largura_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    n = prox(_FIGURA_N)
    legenda(doc, f"Figura {n} — {titulo}. Captura real do sistema em execução "
                 f"({comando_ou_url}).")
    return n


def placeholder_screenshot(doc, titulo, comando_ou_url, esperado, arquivo_sugerido):
    n = prox(_FIGURA_N)
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, "F5F7FB")
    set_cell_borders(cell, color="9BB3D4", sz=10)
    cell.text = ""
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(30)
    p0.paragraph_format.space_after = Pt(4)
    r = p0.add_run(f"[ Figura {n} — espaço reservado para captura de tela ]")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x6C, 0x8D, 0xBA)

    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(30)
    r1 = p1.add_run(f"Salvar como: {arquivo_sugerido}")
    r1.font.size = Pt(9)
    r1.italic = True
    r1.font.color.rgb = CINZA_CLARO

    legenda(doc, f"Figura {n} — {titulo}", space_before=4, space_after=2)
    rich_par(doc, [
        ("Como capturar: ", {"bold": True}), (comando_ou_url, {"italic": True}),
    ], size=9.5, space_after=2)
    rich_par(doc, [
        ("Resultado esperado: ", {"bold": True}), (esperado, {}),
    ], size=9.5, space_after=18)
    return n


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C4D2E7")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ===========================================================================
# CONTEUDO
# ===========================================================================

def secao_capa(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ECOSSISTEMA EnadeX")
    r.font.size = Pt(13)
    r.font.color.rgb = CINZA_CLARO
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Projeto 4 — ENADE-Time Distribuído")
    r.font.size = Pt(15)
    r.font.color.rgb = AZUL_MEDIO

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("Sistema Paralelo de Análise Longitudinal")
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = AZUL_ESCURO

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("dos Microdados do ENADE")
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = AZUL_ESCURO

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento técnico-acadêmico do subprojeto")
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = CINZA_TEXTO

    for _ in range(6):
        doc.add_paragraph()

    campos = [
        ("Aluno responsável", "Lucas Eduardo Tavares Costa"),
        ("Disciplina", "Sistemas Paralelos e Distribuídos (SPD)"),
        ("Critério / Módulos do subprojeto",
         "Critério B — Módulo 1 (ETL Longitudinal e Harmonização) "
         "+ Módulo 3 (Dashboard e Exportação)"),
        ("Diretório do projeto", "C:\\Projetos\\ENADE"),
        ("Data do documento", data_extenso()),
    ]
    for label, valor in campos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  ")
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(valor)
        r2.font.size = Pt(11)
    doc.add_page_break()


def secao_resumo(doc):
    p = doc.add_heading(level=1)
    p.paragraph_format.page_break_before = False
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Resumo")
    r.font.color.rgb = AZUL_ESCURO

    par(doc,
        "O ENADE-Time Distribuído é o subprojeto responsável, dentro do ecossistema "
        "EnadeX, por transformar os microdados brutos do ENADE — arquivos oficiais do "
        "INEP, publicados a cada edição em formatos e esquemas que mudam de ciclo para "
        "ciclo — em uma base longitudinal única, harmonizada e validada, cobrindo os "
        "cursos de Computação do Norte e Nordeste brasileiros entre 2005 e 2021. Sobre "
        "essa base, o subprojeto conduz um experimento de Sistemas Paralelos e "
        "Distribuídos: compara a execução sequencial do pipeline de ETL com execuções "
        "paralelas via multiprocessamento em uma campanha com repetições e duas ordens "
        "de submissão, medindo speedup, eficiência e throughput pareados por suíte e "
        "decompondo a perda observada em granularidade + escalonamento, contenção e "
        "overhead — com a máquina (núcleos físicos e lógicos) e a condição de cache "
        "declaradas. O resultado é entregue em três "
        "camadas integradas — um banco PostgreSQL relacional, uma API FastAPI "
        "somente leitura e um dashboard React interativo — publicadas de forma "
        "reprodutível via Docker, e disponibilizadas ao restante do ecossistema EnadeX "
        "tanto como arquivo CSV consolidado quanto como contrato de API REST "
        "documentado.")

    rich_par(doc, [
        ("Palavras-chave: ", {"bold": True}),
        ("ENADE; microdados educacionais; ETL; harmonização de dados longitudinais; "
         "processamento paralelo; multiprocessing; Lei de Amdahl; PostgreSQL; "
         "FastAPI; React; dashboard analítico.", {"italic": True}),
    ], size=10.5)


def secao_1_problema(doc):
    h1(doc, "1. Problema de pesquisa")

    par(doc,
        "O ENADE (Exame Nacional de Desempenho de Estudantes) é aplicado pelo INEP a "
        "cada três anos e produz, a cada edição, milhões de linhas de microdados "
        "brutos por curso avaliado — arquivos de cadastro (arq1) e de notas (arq3), "
        "entre outros. Tomados isoladamente, esses arquivos já são custosos de "
        "processar: exigem detecção de encoding e separador, filtragem por área e "
        "região, tratamento de sentinelas de valor ausente e agregação de notas por "
        "curso. Tomados longitudinalmente — como exige qualquer análise de evolução "
        "ao longo de ciclos —, o problema se agrava porque o próprio esquema dos "
        "arquivos muda entre edições: colunas aparecem ou desaparecem (é o caso de "
        "CO_MODALIDADE, ausente em 2005 e 2008), e códigos de classificação são "
        "substituídos no meio da série histórica (é o caso do código da área de "
        "Computação, que migra de 40 para 4004 a partir de 2011). Nenhuma dessas "
        "mudanças é documentada em um único lugar pelo INEP; harmonizá-las é trabalho "
        "manual e sujeito a erro se feito ano a ano.")

    par(doc,
        "Definido esse pipeline de harmonização, surge a pergunta que ancora a "
        "vertente de Sistemas Paralelos e Distribuídos do subprojeto: como a unidade "
        "natural de trabalho desse ETL — um ano/edição do ENADE, processado de forma "
        "independente dos demais — se comporta sob paralelismo por processos? Mais "
        "precisamente:")

    bullets(doc, [
        "É possível reduzir de forma mensurável o tempo total do pipeline "
        "distribuindo o processamento dos seis anos do recorte entre múltiplos "
        "processos do sistema operacional?",
        "Qual é o ganho real (speedup) e o aproveitamento de cada processo "
        "adicional (eficiência) nesse cenário, que mistura leitura de arquivo texto "
        "(I/O), parsing com pandas (CPU) e escrita em lote no PostgreSQL (I/O de "
        "rede/disco)?",
        "Existe um número ótimo de processos para o hardware disponível, e o que "
        "explica — em termos de Lei de Amdahl, contenção de I/O e overhead de "
        "criação de processos — o ponto em que adicionar mais processos deixa de "
        "compensar?",
    ])

    par(doc,
        "Em síntese, o problema de pesquisa combina uma dimensão de engenharia de "
        "dados (como consolidar, de forma auditável e reprodutível, uma série "
        "histórica de microdados educacionais cujo esquema muda ao longo do tempo) "
        "com uma dimensão de sistemas paralelos e distribuídos (até que ponto — e por "
        "que motivos técnicos específicos — a paralelização por processos acelera "
        "esse tipo de pipeline predominantemente limitado por I/O).")


def secao_2_objetivo_geral(doc):
    h1(doc, "2. Objetivo geral")

    par(doc,
        "Projetar, implementar e avaliar empiricamente um sistema completo que (i) "
        "harmonize e consolide os microdados do ENADE dos cursos de Computação do "
        "Norte e Nordeste brasileiros entre 2005 e 2021 em uma base longitudinal "
        "única e íntegra; (ii) meça, por meio de um experimento controlado de "
        "processamento sequencial versus paralelo, o impacto real do "
        "multiprocessamento sobre esse pipeline de ETL, explicando os resultados à "
        "luz da Lei de Amdahl e das características de I/O do pipeline; e (iii) "
        "disponibilize tanto a base consolidada quanto os resultados do experimento "
        "por meio de uma API somente leitura e de um dashboard interativo, "
        "integrando-se de forma documentada ao ecossistema EnadeX.")


def secao_3_objetivos_especificos(doc):
    h1(doc, "3. Objetivos específicos")

    numbered(doc, [
        ("Harmonizar o esquema entre edições — ",
         "detectar automaticamente encoding e separador de cada arquivo bruto, "
         "tratar a ausência de CO_MODALIDADE em 2005/2008 e reconciliar a mudança de "
         "código da área de Computação (40 → 4004) sem descartar o valor original do "
         "ano, preservando a auditabilidade do dado bruto."),
        ("Consolidar e validar a base longitudinal — ",
         "produzir um único arquivo consolidado (24.967 registros, 13 colunas, seis "
         "edições) com validação automática de escala das notas ([0, 100]), "
         "coerência de nulos entre as três notas (FG, GER, CE) e contagens por ano "
         "auditáveis contra os arquivos brutos."),
        ("Modelar e persistir os dados em um banco relacional — ",
         "projetar um esquema dimensional (tabela fato + sete dimensões) em "
         "PostgreSQL, executado em container Docker para reprodutibilidade, com "
         "carga em lote via COPY FROM STDIN."),
        ("Implementar e instrumentar as duas variantes do pipeline de ETL — ",
         "uma execução sequencial (baseline) e execuções paralelas com "
         "ProcessPoolExecutor (2, 3, 4 e 6 processos, em duas ordens de submissão, "
         "com repetições), registrando tempo total, throughput, memória de pico, "
         "CPU média do sistema e bytes lidos do disco de cada execução em tabelas "
         "dedicadas do banco."),
        ("Calcular e expor as métricas clássicas de paralelismo — ",
         "speedup (S(p) = T(1)/T(p)) e eficiência (E(p) = S(p)/p) em uma única "
         "definição — uma view SQL que pareia cada execução paralela com o "
         "sequencial da PRÓPRIA suíte — e agregar mediana, mínimo, máximo e IQR "
         "por configuração em uma segunda view; a API apenas lê as views."),
        ("Expor os dados e os resultados via API REST somente leitura — ",
         "implementar uma API em FastAPI, sem operações de escrita, organizada em "
         "grupos de endpoints para dimensões, análises agregadas e benchmark, "
         "documentada automaticamente via OpenAPI/Swagger."),
        ("Construir um dashboard analítico interativo — ",
         "um painel em React que apresente as análises longitudinais do ENADE "
         "(por ano, região, UF, IES) e os resultados do experimento de "
         "paralelismo, incluindo comparação entre grupos, alerta de amostra "
         "reduzida e exportação de dados e relatórios."),
        ("Documentar e disponibilizar o contrato de integração com o ecossistema "
         "EnadeX — ",
         "publicar a base consolidada em formato consumível por outros "
         "subprojetos (CSV direto ou API REST), com o esquema, as garantias de "
         "imutabilidade e o versionamento formalmente documentados."),
    ])


def secao_4_justificativa(doc):
    h1(doc, "4. Justificativa do projeto")

    h2(doc, "4.1 Relevância acadêmica para Sistemas Paralelos e Distribuídos")
    par(doc,
        "O pipeline de ETL do ENADE é um caso de estudo particularmente honesto para "
        "SPD: não é um problema artificialmente paralelizável (como multiplicação de "
        "matrizes ou soma de vetores), mas um pipeline real, misto de I/O e CPU, "
        "exatamente o tipo de carga em que a Lei de Amdahl e o overhead de "
        "paralelização se manifestam de forma visível. Medir speedup e eficiência "
        "neste cenário — e observar que o ponto ótimo depende do número de núcleos "
        "físicos, da ordem de submissão das tarefas e da granularidade do trabalho — "
        "é mais instrutivo do que confirmar paralelismo ideal em uma carga sintética.")

    h2(doc, "4.2 Relevância para o ecossistema EnadeX e para a área de Computação")
    par(doc,
        "O recorte escolhido (Computação, Norte e Nordeste, seis edições trienais) "
        "cobre uma lacuna real: dados oficiais e harmonizados sobre a evolução do "
        "desempenho de cursos de Computação nas regiões historicamente menos "
        "representadas nas análises do ENADE. Ao disponibilizar essa base já limpa, "
        "validada e documentada — em vez de apenas os arquivos brutos do INEP — o "
        "subprojeto reduz o custo de entrada para qualquer outro módulo do "
        "ecossistema EnadeX que precise da mesma série histórica.")

    h2(doc, "4.3 Rigor de engenharia como pré-condição para a pesquisa")
    par(doc,
        "Um experimento de desempenho só é confiável se a base sobre a qual ele roda "
        "for íntegra. Por isso o projeto trata harmonização e validação de dados "
        "como parte do mesmo esforço que o benchmark de paralelismo: a mesma "
        "disciplina de reprodutibilidade (Docker, scripts idempotentes, validação "
        "linha a linha) que sustenta a confiabilidade da base também sustenta a "
        "confiabilidade das medições de tempo. O projeto documentou, inclusive, um "
        "incidente real de medição — execuções de teste feitas com o diretório "
        "sincronizado pelo OneDrive geraram contenção de I/O e precisaram ser "
        "descartadas do conjunto oficial — o que reforça, na prática, por que "
        "ambiente controlado é parte do método, não um detalhe operacional.")

    h2(doc, "4.4 Reprodutibilidade e auditabilidade como requisitos de projeto")
    par(doc,
        "Todas as decisões de maior impacto — a equivalência entre os códigos 40 e "
        "4004, o tratamento de CO_MODALIDADE ausente, a métrica de speedup, os "
        "identificadores de execução de benchmark excluídos por padrão — são "
        "calculadas em SQL auditável ou documentadas explicitamente, nunca ocultas "
        "em lógica de aplicação difícil de inspecionar. Essa escolha foi orientada "
        "pelo padrão acadêmico de que qualquer resultado apresentado precisa ser "
        "reproduzível e verificável por terceiros, incluindo a banca avaliadora.")


def secao_5_ferramentas(doc):
    h1(doc, "5. Ferramentas, tecnologias, linguagens, bibliotecas e frameworks")

    par(doc,
        "A tabela a seguir relaciona o ferramental efetivamente utilizado na "
        "implementação, agrupado por categoria e com o papel específico de cada "
        "item dentro do sistema.")

    headers = ["Categoria", "Item", "Versão", "Papel no projeto"]
    rows = [
        ("Linguagem", "Python", "3.13", "ETL, harmonização, benchmark paralelo, API"),
        ("Linguagem", "TypeScript", "5.6", "Tipagem estática do frontend, espelha os schemas Pydantic"),
        ("Framework — API", "FastAPI", "≥ 0.115", "API REST somente leitura, OpenAPI/Swagger automático"),
        ("Validação de dados", "Pydantic v2 / pydantic-settings", "≥ 2.7 / ≥ 2.4", "Contratos HTTP e configuração via variáveis de ambiente"),
        ("Driver de banco", "psycopg2-binary", "≥ 2.9", "Acesso ao PostgreSQL sem ORM; ThreadedConnectionPool"),
        ("Servidor ASGI", "Uvicorn", "≥ 0.30", "Execução da aplicação FastAPI"),
        ("Banco de dados", "PostgreSQL", "16 (local) / 17.6 (Supabase)", "Armazenamento relacional dimensional"),
        ("Orquestração", "Docker / Docker Compose", "—", "Reprodutibilidade do banco e (opcionalmente) da API"),
        ("Paralelismo", "concurrent.futures.ProcessPoolExecutor", "biblioteca padrão", "Núcleo do experimento de SPD — paralelismo por processos"),
        ("Monitoramento de recursos", "psutil", "—", "Coleta de CPU, memória e metadados de máquina no benchmark"),
        ("Manipulação de dados", "pandas", "3.0", "Leitura, harmonização, agregação e preparação de buffers CSV"),
        ("Computação numérica", "NumPy", "—", "Suporte a pandas nas conversões e sentinelas de nulo"),
        ("Framework — frontend", "React", "18.3", "Interface do dashboard (SPA)"),
        ("Build / dev server", "Vite", "5.4", "Bundler e servidor de desenvolvimento do frontend"),
        ("Roteamento", "React Router", "6.27", "Navegação da SPA; filtros persistidos como deep-link na URL"),
        ("Estado de servidor", "TanStack Query", "5.59", "Cache, loading, erro e refetch das chamadas à API"),
        ("Cliente HTTP", "Axios", "1.7", "Chamadas HTTP do frontend para a API"),
        ("Visualização", "Recharts", "2.13", "Gráficos de linha, barra e comparação"),
        ("Estilo", "Tailwind CSS", "3.4", "Design system utilitário do dashboard"),
        ("Qualidade de código", "ESLint + typescript-eslint", "9.x", "Lint estático do frontend"),
        ("Plataforma em nuvem (opcional)", "Supabase", "—", "Postgres gerenciado; réplica pública de leitura da base consolidada"),
        ("Controle de versão", "Git / GitHub", "—", "Versionamento do código-fonte"),
    ]
    tabela(doc, headers, rows, widths_cm=[3.2, 4.6, 2.6, 6.2],
           legenda_texto="Ferramentas, tecnologias e bibliotecas utilizadas, por categoria.",
           font_size=9.5)


def secao_6_processo(doc):
    h1(doc, "6. Processo de desenvolvimento e implantação")

    h2(doc, "6.1 Desenvolvimento — decisões e etapas, em ordem cronológica")
    par(doc,
        "O desenvolvimento seguiu um registro formal de decisões (design log), do "
        "qual as etapas abaixo são a síntese ordenada.")

    numbered(doc, [
        ("Definição do recorte do estudo — ",
         "decidiu-se processar apenas os cursos de Computação (CO_GRUPO 40/4004) "
         "nas regiões Norte e Nordeste (CO_REGIAO_CURSO 1/2), nas seis edições "
         "trienais disponíveis (2005, 2008, 2011, 2014, 2017, 2021), para manter o "
         "experimento tratável sem perder representatividade longitudinal."),
        ("Inspeção da estrutura dos arquivos brutos — ",
         "script dedicado varre cada pasta de edição, detecta encoding "
         "(utf-8/latin1/cp1252) e separador (;/,/tab) de cada arquivo, e registra a "
         "presença ou ausência de cada coluna relevante por ano."),
        ("Implementação da harmonização — ",
         "conversão de colunas inteiras com tratamento de sentinelas de valor "
         "ausente, normalização de decimal (vírgula → ponto) nas notas, criação de "
         "CO_MODALIDADE com nulo quando a coluna não existe no arquivo bruto do ano, "
         "e filtragem pelo recorte definido na etapa 1."),
        ("Consolidação e primeira validação — ",
         "agregação das notas por curso (média de NT_FG por CO_CURSO), junção com o "
         "cadastro filtrado e concatenação de todos os anos em um único CSV, com "
         "validação de escala das notas e contagem de linhas por ano."),
        ("Modelagem do esquema relacional — ",
         "projeto de um esquema dimensional (tabela fato_enade + sete dimensões: "
         "região, UF, grupo, categoria administrativa, organização acadêmica, "
         "modalidade, ano), com chaves estrangeiras, checks de escala [0,100] e "
         "índices compostos para os filtros mais usados pela API."),
        ("Carga do banco via COPY FROM STDIN — ",
         "carga em lote do CSV consolidado no PostgreSQL, com correção específica "
         "para colunas inteiras que chegavam como float por causa de valores nulos "
         "(bug de formatação do tipo \"2806701.000000\"), resolvida com coerção "
         "para o tipo nullable Int64 do pandas antes da escrita do buffer CSV."),
        ("Implementação do pipeline sequencial de referência — ",
         "versão single-thread do processamento dos seis anos, usada como "
         "baseline T(1) para todas as métricas de speedup e eficiência."),
        ("Implementação do pipeline paralelo — ",
         "extração da lógica de processamento de um ano para uma função pura e "
         "importável no nível do módulo (exigência do modelo spawn do "
         "multiprocessing no Windows), paralelizada com ProcessPoolExecutor, com "
         "granularidade de um processo por ano."),
        ("Instrumentação e persistência das métricas de benchmark — ",
         "cada execução (sequencial ou paralela) grava tempo total, throughput, "
         "memória de pico e metadados de máquina em benchmark_execucao, com o "
         "detalhamento por ano em benchmark_etapa, em uma única transação."),
        ("Cálculo de speedup e eficiência — ",
         "implementação de uma view SQL (v_benchmark_metricas) que pareia cada "
         "execução paralela com o sequencial da própria suíte e calcula "
         "S(p) = T(1)/T(p) e E(p) = S(p)/p diretamente em SQL, auditável "
         "independentemente do código Python; linhas antigas, sem suíte, mantêm o "
         "pareamento temporal com o sequencial imediatamente anterior."),
        ("Execução da suíte oficial de benchmark — ",
         "execução sequencial e paralela (2 e 4 processos) fora de qualquer "
         "diretório sincronizado por serviços de nuvem, após identificar que a "
         "sincronização do OneDrive contaminava as medições de I/O de uma rodada "
         "de teste anterior."),
        ("Construção da API somente leitura — ",
         "roteadores separados por domínio (saúde, dimensões, análises, "
         "benchmark), pool de conexões compartilhado, CORS restrito às origens do "
         "frontend local, e nenhuma rota de escrita."),
        ("Construção do dashboard — ",
         "páginas para as análises do ENADE (anual, regional, por UF, ranking de "
         "IES, comparação entre grupos, registros paginados) e para o benchmark "
         "SPD (execuções, comparativo, métricas, detalhamento por etapa), com "
         "filtros refletidos na URL e exportação de dados/relatórios."),
        ("Evolução controlada do schema de dados (v1.0 → v1.1) — ",
         "acréscimo das notas MEDIA_NT_GER e MEDIA_NT_CE ao consolidado por meio de "
         "um utilitário que reconstrói os dados a partir dos arquivos brutos, "
         "valida linha a linha contra a base anterior e só grava se a validação "
         "for 100% bem-sucedida, com backup automático prévio."),
        ("Integração com o ecossistema EnadeX — ",
         "documentação do contrato de consumo da base (CSV direto ou API REST) "
         "para os demais subprojetos, e replicação opcional da tabela consolidada "
         "em um projeto Supabase na nuvem."),
        ("Correção do baseline do comparativo — ",
         "identificou-se que o endpoint /api/benchmark/comparativo dividia todas "
         "as execuções pelo sequencial mais recente, misturando rodadas; a API "
         "passou a ler exclusivamente a view (pareamento por suíte), com checagem "
         "cruzada em Python (scripts/13_validar_metricas.py) e testes automatizados."),
        ("Ordem de submissão como variável do experimento — ",
         "o ProcessPoolExecutor entrega cada tarefa ao primeiro worker livre na "
         "ordem em que foi submetida; o script paralelo ganhou a opção LPT (Graham, "
         "1969), usando os tempos por ano do sequencial da mesma suíte, e a ordem "
         "efetivamente usada é gravada em cada execução."),
        ("Campanha oficial com repetições e instrumentação — ",
         "cinco suítes, cada uma com sequencial + {2, 3, 4, 6} workers × {crescente, "
         "LPT}; aquecimento de cache descartado; CPU média e bytes lidos do disco "
         "por execução; medianas, mínimos, máximos e IQR em v_benchmark_resumo."),
    ])

    h2(doc, "6.2 Implantação — como o sistema é colocado em execução")
    par(doc,
        "A implantação local segue quatro etapas independentes, documentadas em "
        "detalhe no guia de execução do repositório.")

    numbered(doc, [
        ("Banco de dados — ",
         "docker compose up -d postgres sobe um container PostgreSQL 16 com "
         "healthcheck; scripts dedicados criam o esquema e carregam o CSV "
         "consolidado na primeira execução."),
        ("Validação do banco — ",
         "um script de validação confirma a contagem de 24.967 linhas, a "
         "integridade referencial contra as sete dimensões, a coerência dos "
         "nulos entre as três notas e a presença de todas as tabelas esperadas."),
        ("API — ",
         "executada no host via Uvicorn com recarregamento automático durante o "
         "desenvolvimento, ou em container por meio de um serviço opcional do "
         "docker-compose (perfil \"api\"), que também conecta ao mesmo banco."),
        ("Frontend — ",
         "npm run dev inicia o servidor de desenvolvimento Vite; npm run build "
         "gera os artefatos estáticos de produção, servidos localmente com "
         "npm run preview."),
    ])


def secao_7_arquitetura_geral(doc):
    h1(doc, "7. Arquitetura geral do sistema")

    par(doc,
        "O sistema é organizado em seis camadas verticais, cada uma dependendo "
        "apenas da anterior, complementadas por dois consumidores externos da base "
        "consolidada (Figura 1).")

    imagem(doc, DIAGRAMAS / "arquitetura_geral.png",
           "Arquitetura geral do ENADE-Time Distribuído — pipeline principal em seis "
           "camadas, com os ramos de benchmark paralelo, réplica em nuvem e consumo "
           "pelo ecossistema EnadeX.")

    h2(doc, "7.1 Componentes")
    bullets(doc, [
        ("Dados brutos (INEP): arquivos texto oficiais do ENADE por edição, já "
         "anonimizados sob a LGPD pelo próprio INEP, tratados como somente leitura "
         "em todo o pipeline — nenhum script grava nesta camada.",
         "Camada 1 — "),
        ("ETL e harmonização: detecta encoding/separador, resolve as diferenças de "
         "esquema entre edições (coluna ausente, código de área migrado), filtra "
         "pelo recorte do estudo e converte notas e sentinelas de nulo.",
         "Camada 2 — "),
        ("Consolidado (CSV): artefato intermediário estável — 24.967 linhas, 13 "
         "colunas — que serve tanto de entrada para a carga no PostgreSQL quanto "
         "para a réplica opcional no Supabase.",
         "Camada 3 — "),
        ("PostgreSQL (Docker): banco relacional dimensional com a tabela fato, as "
         "sete dimensões, as tabelas de benchmark e as views que calculam speedup e "
         "eficiência por suíte e os agregados por campanha; roda em container para "
         "garantir que qualquer máquina "
         "reproduza o mesmo ambiente.",
         "Camada 4 — "),
        ("API FastAPI: única porta de entrada HTTP para o dado, estritamente "
         "somente leitura (sem POST/PUT/DELETE), com CORS restrito às origens "
         "locais do frontend.",
         "Camada 5 — "),
        ("Dashboard React: única interface do usuário final; nunca acessa o banco "
         "diretamente, sempre por meio da API; mantém o estado dos filtros na URL.",
         "Camada 6 — "),
    ])

    h2(doc, "7.2 Fluxo de dados e principais decisões técnicas")
    par(doc,
        "O fluxo principal é estritamente unidirecional — dado bruto → CSV "
        "consolidado → banco → API → dashboard — o que simplifica a depuração, "
        "pois qualquer divergência observada no dashboard pode ser rastreada "
        "camada a camada até a origem. Três decisões técnicas sustentam essa "
        "arquitetura:")
    bullets(doc, [
        ("PostgreSQL em vez de SQLite ou DuckDB — ",
         "o experimento de paralelismo depende de um servidor real capaz de "
         "aceitar múltiplas conexões simultâneas dos processos de benchmark; "
         "SQLite não atenderia a esse requisito, e DuckDB deslocaria o foco do "
         "componente \"servidor de banco\" que a disciplina de SPD exige."),
        ("API sem ORM — ",
         "o domínio é tabular e somente leitura; SQL explícito, parametrizado, é "
         "mais auditável para um projeto acadêmico do que a camada adicional de "
         "abstração de um ORM."),
        ("Estado do filtro na URL, não em um gerenciador de estado global — ",
         "cada combinação de filtros do dashboard vira um link direto "
         "(deep-link), o que facilita tanto o compartilhamento de uma visão "
         "específica quanto a depuração durante a apresentação do projeto."),
    ])


def secao_8_integracao_subprojeto(doc):
    h1(doc, "8. Integração do subprojeto com o sistema maior")

    par(doc,
        "O ENADE-Time Distribuído é o Projeto 4 do ecossistema EnadeX, respondendo "
        "ao Critério B do currículo do curso. Dentro desse critério, o subprojeto "
        "entrega dois dos três módulos previstos: o Módulo 1 (ETL Longitudinal e "
        "Harmonização) e o Módulo 3 (Dashboard e Exportação). O Módulo 2 — que "
        "envolveria estatística inferencial mais pesada, como intervalos de "
        "confiança de 95%, regressão/LOESS e detecção de quebra estrutural na "
        "série temporal — foi conscientemente deixado como complementar, por "
        "exigir um escopo estatístico distinto do que este subprojeto se propôs a "
        "cobrir; essa decisão está registrada no histórico de decisões do projeto "
        "e não é apresentada aqui como concluída.")

    h2(doc, "8.1 O que este subprojeto contribui para o ecossistema")
    par(doc,
        "O ENADE-Time Distribuído é, ao mesmo tempo, consumidor dos microdados "
        "brutos do INEP e provedor de dados para o restante do ecossistema EnadeX. "
        "O contrato de consumo — o que a base oferece, em que formato, com que "
        "garantias — está documentado internamente e resumido a seguir:")

    headers = ["Item", "Valor"]
    rows = [
        ("Artefato principal", "CSV consolidado (24.967 linhas, 13 colunas, seis edições)"),
        ("Formas de consumo", "Arquivo CSV direto (offline) ou API REST somente leitura"),
        ("Endpoint-base da API", "http://localhost:8000/api (Swagger em /docs)"),
        ("Filtros combináveis", "ano, região, UF, IES, grupo, modalidade, categoria administrativa, organização acadêmica"),
        ("Garantia de imutabilidade", "colunas e grão fixos dentro de uma versão; mudança de esquema gera nova versão"),
        ("Tratamento de nulos", "notas ausentes permanecem NULL — nunca são imputadas"),
        ("Versão atual", "v1.1 (acréscimo controlado de MEDIA_NT_GER e MEDIA_NT_CE sobre a v1.0)"),
    ]
    tabela(doc, headers, rows, widths_cm=[4.5, 12.1],
           legenda_texto="Contrato de consumo da base oferecido pelo subprojeto ao ecossistema EnadeX.")

    par(doc,
        "Vale registrar, com a mesma transparência adotada na documentação interna "
        "do projeto, que este contrato descreve o que o subprojeto disponibiliza — "
        "não uma integração já formalizada com outro módulo específico do "
        "ecossistema. Até a data deste documento, nenhum acordo conjunto de "
        "consumo foi combinado e efetivado com outro subprojeto; o espaço para "
        "registrar esse acordo (projeto consumidor, formato acordado, data) está "
        "preparado na documentação interna e será preenchido quando essa decisão "
        "existir de fato.")

    h2(doc, "8.2 Arquitetura específica do subprojeto — o benchmark paralelo")
    par(doc,
        "A contribuição mais específica deste subprojeto para a disciplina de SPD "
        "é o mecanismo de benchmark paralelo, detalhado na Figura 2. Um processo "
        "principal distribui os seis anos do recorte entre N processos "
        "trabalhadores (N = 1 no baseline sequencial; N = 2, 3, 4 ou 6 nas execuções "
        "paralelas, em duas ordens de submissão) usando ProcessPoolExecutor. Cada "
        "trabalhador executa uma "
        "função pura — sem escrita em disco e sem conexão de banco — que "
        "processa um ano inteiro de forma independente. Somente o processo "
        "principal, após reunir os resultados de todos os trabalhadores, grava as "
        "métricas no PostgreSQL em uma única transação. A partir daí, a mesma "
        "métrica de speedup calculada em SQL alimenta tanto a API quanto o "
        "dashboard, garantindo que o número mostrado na tela seja idêntico ao "
        "número auditável no banco.")

    imagem(doc, DIAGRAMAS / "arquitetura_subprojeto.png",
           "Mecânica interna do benchmark paralelo — do processo principal aos "
           "processos trabalhadores, até a métrica exibida no dashboard.")

    par(doc,
        "Essa separação — trabalhadores puros e sem I/O de banco, processo "
        "principal como único ponto de escrita — não é um detalhe de "
        "implementação: é a resposta a uma restrição real do Windows, onde o "
        "multiprocessing usa o método spawn (cada processo novo reimporta o "
        "módulo Python do zero, em vez de herdar a memória do processo pai via "
        "fork, como ocorre no Linux). Isso exige que a função executada pelos "
        "trabalhadores seja importável no nível do módulo e livre de efeitos "
        "colaterais — decisão de projeto explicada em detalhe no Quadro 1, na "
        "Seção 10.")


def secao_8_3_resultados(doc):
    """Resultados da campanha oficial — tudo vem de D (banco), nada digitado."""
    maq = D["maquina"]
    esc = D["escalonamento"]
    seq = D["sequencial"]
    melhor = D["melhor"]
    n_suites = maq["n_suites"]
    _parciais = n_suites - maq.get("n_suites_completas", n_suites)
    suites_frase = f"{maq.get('n_suites_completas', n_suites)} suítes completas"
    if _parciais:
        suites_frase += (f" e {_parciais} parcial (interrompida por desligamento da máquina; as execuções "
                         f"gravadas permanecem válidas e elevam o n de algumas configurações)")
    hip = esc["hipoteses"]
    disco_aq = D["aquecimento"][0]["disco_mb"] if D["aquecimento"] else None

    h2(doc, "8.3 Resultados do experimento — campanha oficial")
    par(doc,
        f"A campanha oficial (identificador {D['campanha_id']}) foi executada entre "
        f"{maq['inicio'].strftime('%d/%m/%Y %H:%M')} e {maq['fim'].strftime('%d/%m/%Y %H:%M')} UTC, "
        f"com {suites_frase}. Cada suíte completa "
        f"roda um sequencial e, em seguida, {len(D['workers'])} tamanhos de pool "
        f"({', '.join(str(w) for w in D['workers'])} workers) em duas ordens de "
        f"submissão (crescente e LPT), totalizando {maq['n_execucoes']} execuções "
        f"oficiais. Cada execução paralela é pareada com o sequencial da própria "
        f"suíte; os agregados abaixo são medianas sobre as suítes da campanha (o n de cada configuração está nas tabelas), com "
        f"mínimo e máximo.")

    h3(doc, "Máquina e condições declaradas")
    bullets(doc, [
        (f"{maq['cpu_modelo']} — {maq['cpu_fisicos']} núcleos físicos / "
         f"{maq['cpu_logicos']} lógicos (Intel Core i5-1135G7), NVMe, 24 GB de RAM, Windows 11. "
         f"Consequência: p = 4 satura os núcleos físicos; p = 6 roda em hyperthreads.", "Processador: "),
        (f"passada de aquecimento descartada antes da campanha (leu {br(disco_aq, 0)} MB do disco); "
         f"todas as execuções oficiais rodaram com cache de páginas quente — "
         f"mediana de {br(hip.get('H2_disco_quente_mb_mediana'), 1)} MB lidos por execução.", "Cache: "),
        (f"CPU ociosa medida por 10 s antes da campanha: {br(maq['cpu_ocioso'], 1)} % "
         f"(Docker Desktop com o container do PostgreSQL ativo); CPU média do sistema durante as "
         f"execuções: {br(maq['cpu_percent_medio'], 1)} %.", "Carga de fundo: "),
        ("cada execução grava tempo, throughput, memória de pico, CPU média, bytes lidos, "
         "ordem de submissão, suíte e campanha; os workers continuam puros (sem I/O nem banco).",
         "Instrumentação: "),
    ])

    h3(doc, "Resultados por configuração")
    rows = []
    for r in D["resumo"]:
        rows.append((
            r["rotulo"], str(r["n"]),
            f"{br(r['tempo_mediana'])} [{br(r['tempo_min'])}–{br(r['tempo_max'])}]",
            br(r["tempo_iqr"]),
            sx(r["speedup_mediana"]) if r["modo"] == "paralelo" else "1,0000×",
            (f"{sx(r['speedup_min'], 2)}–{sx(r['speedup_max'], 2)}" if r["modo"] == "paralelo" else "—"),
            pct(r["eficiencia_mediana"]) if r["modo"] == "paralelo" else "100,0 %",
            br(r["throughput_mediana"], 0),
        ))
    tabela(doc, ["Configuração", "n", "Tempo (s) mediana [mín–máx]", "IQR (s)", "Speedup (mediana)",
                 "Speedup mín–máx", "Eficiência", "Throughput (linhas/s)"],
           rows, widths_cm=[3.4, 0.8, 3.3, 1.3, 2.0, 2.2, 1.8, 2.2],
           legenda_texto=f"Campanha oficial — mediana, mínimo, máximo e IQR sobre as suítes da campanha, por configuração (n na tabela) "
                         f"(fonte: v_benchmark_resumo).", font_size=9)
    if melhor:
        par(doc,
            f"A melhor configuração foi {melhor['rotulo']}: speedup mediano de "
            f"{sx(melhor['speedup_mediana'])} (mín–máx {sx(melhor['speedup_min'], 2)}–"
            f"{sx(melhor['speedup_max'], 2)}) e eficiência de {pct(melhor['eficiencia_mediana'])}, "
            f"contra um sequencial mediano de {seg(seq['tempo_mediana'])}.")

    h3(doc, "Tetos de escalonamento e decomposição da perda")
    par(doc,
        "O ProcessPoolExecutor entrega cada ano ao primeiro worker livre, na ordem em que "
        "foi submetido — não faz balanceamento. Com seis unidades de tamanhos distintos, "
        "isso fixa um teto de speedup que independe de Amdahl: o makespan do "
        "escalonamento guloso na ordem usada, calculado com os tempos por ano do "
        "sequencial da mesma suíte. A tabela abaixo separa três perdas: do ideal ao teto "
        "(granularidade + escalonamento), do teto ao makespan medido (contenção: as "
        "etapas ficam mais lentas em paralelo) e do makespan ao wall-clock (overhead de "
        "spawn, coleta e gravação).")
    rows = []
    for r in esc["agregado"]:
        rows.append((
            str(r["workers"]), "LPT" if r["ordem"] == "lpt" else "crescente", str(r["n"]),
            br(r["ideal"]), f"{br(r['teto_real'])} → {sx(r['speedup_teto_real'], 2)}",
            f"{br(r['teto_lpt'])} → {sx(r['speedup_teto_lpt'], 2)}",
            br(r["makespan_medido"]), f"{br(r['inflacao_etapas'], 2)}×", br(r["wall"]), br(r["overhead"], 1),
            sx(r["speedup_medido"], 4), pct(r["pct_do_teto_real"], 0),
        ))
    tabela(doc, ["p", "Ordem", "n", "Ideal (s)", "Teto ordem usada → S", "Teto LPT → S",
                 "Makespan medido (s)", "Inflação", "Wall (s)", "Overhead (s)", "S medido", "% do teto"],
           rows, widths_cm=[0.6, 1.5, 0.6, 1.2, 2.2, 2.0, 1.6, 1.2, 1.2, 1.3, 1.5, 1.3],
           legenda_texto="Tetos por granularidade + escalonamento (ordem usada e LPT), makespan medido, "
                         "inflação das etapas, overhead e fração do teto atingida — medianas por configuração "
                         "(fonte: benchmark_etapa e v_benchmark_metricas, via analise_escalonamento.py).",
           font_size=8.5)

    h3(doc, "Hipóteses pré-registradas")
    itens = []
    for o in ("lpt", "crescente"):
        h = hip.get(f"H1_p6_nao_supera_p4_{o}")
        if h:
            itens.append((f"speedup mediano p = 4: {sx(h['speedup_p4'])}; p = 6: {sx(h['speedup_p6'])} → "
                          f"{'confirmada' if h['confirmada'] else 'refutada'} ({'LPT' if o == 'lpt' else 'crescente'}). "
                          f"Com 4 núcleos físicos, p = 6 roda em hyperthreads e o teto de granularidade "
                          f"(6 unidades) não cresce a partir de 6.", "H1 — p = 6 não supera p = 4: "))
    itens.append((f"mediana de {br(hip.get('H2_disco_quente_mb_mediana'), 1)} MB lidos do disco nas execuções "
                  f"quentes, contra {br(disco_aq, 0)} MB no aquecimento: o pipeline em regime quente é limitado "
                  f"por CPU (parsing com pandas), não por disco. A explicação anterior — '4 workers lendo do "
                  f"mesmo disco geram contenção' — não se sustenta nos bytes medidos.",
                  "H2 — cache quente, limite de CPU: "))
    h3s = []
    for p in (2, 3, 4):
        h = hip.get(f"H3_lpt_ge_crescente_p{p}")
        if h:
            h3s.append(f"p = {p}: LPT {sx(h['lpt'])} × crescente {sx(h['crescente'])} "
                       f"({'confirmada' if h['confirmada'] else 'refutada'})")
    if h3s:
        itens.append(("; ".join(h3s) + ". A diferença entre as duas ordens é perda de escalonamento — "
                      "o código dos workers é o mesmo.", "H3 — LPT ≥ crescente: "))
    bullets(doc, itens)

    h3(doc, "Histórico das medições e o erro do baseline")
    par(doc,
        "O banco preserva as rodadas anteriores (nenhuma linha foi apagada). A versão "
        "anterior do endpoint de comparativo dividia todas as execuções pelo sequencial "
        "mais recente; quando o banco passou a conter duas rodadas, isso produziu "
        "speedups inflados. A tabela abaixo mostra o efeito, calculado a partir das "
        "próprias linhas do banco:")
    rows = []
    for b in D["bug_baseline"]:
        rows.append((f"#{b['exec']} ({b['workers']} workers, {seg(b['tempo'])})",
                     f"#{b['baseline_errado']} ({seg(b['t_errado'])})", sx(b["s_errado"]), pct(b["e_errado"]),
                     f"#{b['baseline_certo']}", sx(b["s_certo"]), pct(b["e_certo"])))
    if rows:
        tabela(doc, ["Execução", "Baseline usado (errado)", "Speedup exibido", "Eficiência exibida",
                     "Baseline pareado", "Speedup correto", "Eficiência correta"],
               rows, widths_cm=[3.4, 3.0, 2.0, 2.0, 1.9, 2.0, 2.2],
               legenda_texto="Efeito do baseline errado nas rodadas históricas (fonte: benchmark_execucao e v_benchmark_metricas).",
               font_size=9)
    rows = []
    for rod in D["historico"]:
        b = rod["baseline"]
        for p in rod["paralelas"]:
            rows.append((b["timestamp_inicio"].strftime("%d/%m/%Y"), f"#{b['id']}", seg(b["tempo_total_seg"]),
                         f"#{p['id']} ({p['num_workers']} workers)", seg(p["tempo_total_seg"]),
                         sx(p["speedup"]), pct(p["eficiencia"])))
    if rows:
        tabela(doc, ["Data", "Sequencial", "T(1)", "Paralela", "T(p)", "Speedup", "Eficiência"],
               rows, widths_cm=[2.2, 1.8, 2.0, 3.4, 2.0, 2.4, 2.4],
               legenda_texto="Rodadas históricas com pareamento temporal correto (ordem de submissão crescente).",
               font_size=9)
    par(doc,
        "A rodada citada na apresentação anterior (sequencial de aproximadamente 243 s) foi perdida em um "
        "reset do schema e não tem suporte em banco; por isso não é mais citada como "
        "resultado. Os números oficiais deste documento são exclusivamente os da campanha "
        "acima, e um script (docs/geradores/verificar_numeros.py) confere que cada número "
        "citado aqui existe nas views do banco.")


def secao_9_capturas(doc):
    h1(doc, "9. Capturas de tela das principais interfaces e funcionalidades")

    par(doc,
        "As figuras desta seção seguem o checklist de evidências mantido no "
        "repositório do projeto (docs/EVIDENCIAS_TESTES.md). Todas retratam o "
        "sistema realmente em execução: as capturas das interfaces web — API e "
        "dashboard — foram obtidas por navegador Chrome em modo headless "
        "controlado por Puppeteer, com o sistema completo no ar (PostgreSQL em "
        "Docker, API FastAPI e frontend React); as evidências de terminal "
        "reproduzem a saída genuína de cada comando, executado nesta máquina e "
        "tipografada em fonte monoespaçada para legibilidade — nenhum valor foi "
        "editado. As capturas de benchmark (Figuras 5 a 7) e do comparativo "
        "(Figuras 11 e 14) são da sessão de 21/08/2026 e retratam a rodada "
        "histórica daquela data — inclusive o comparativo com o baseline errado "
        "descrito na Seção 8.3. Os números oficiais são os da campanha da Seção "
        "8.3; o dashboard passou a exibir a campanha oficial (mediana por "
        "configuração) após a correção.")

    h2(doc, "9.1 Infraestrutura")
    screenshot_real(
        doc, "Docker Compose — container enade_postgres \"Up (healthy)\" na "
        "porta 5432 e API enade_api em execução (saída real do comando)",
        "screenshots/01_docker_ps.png", "docker compose ps")
    screenshot_real(
        doc, "Validação do banco — 24.967 linhas em fato_enade, integridade "
        "referencial e coerência de nulos confirmadas: \"BASE VALIDADA COM "
        "SUCESSO\" (saída real do comando)",
        "screenshots/02_validar_banco.png",
        "python scripts\\07_validar_banco.py")

    h2(doc, "9.2 Benchmark sequencial × paralelo")
    screenshot_real(
        doc, "Benchmark sequencial (baseline) — execução real com tempo total, "
        "throughput e gravação em benchmark_execucao (saída real do comando; rodada "
        "histórica de 21/08/2026)",
        "screenshots/03_bench_sequencial.png",
        "python scripts\\08_benchmark_sequencial.py")
    screenshot_real(
        doc, "Benchmark paralelo com 2 processos — cada worker processa um ano "
        "inteiro; a saída mostra o PID do worker por etapa (saída real do "
        "comando; rodada histórica de 21/08/2026, ordem crescente)",
        "screenshots/04_bench_2w.png",
        "python scripts\\09_benchmark_paralelo.py --workers 2")
    screenshot_real(
        doc, "Benchmark paralelo com 4 processos — mesmo dataset e máquina "
        "(saída real do comando; rodada histórica de 21/08/2026; os pids por ano "
        "mostram o despacho na ordem de submissão)",
        "screenshots/05_bench_4w.png",
        "python scripts\\09_benchmark_paralelo.py --workers 4")

    h2(doc, "9.3 API FastAPI")
    screenshot_real(
        doc, "Swagger UI — documentação automática da API, com os quatro grupos "
        "de endpoints (health, dimensoes, analises, benchmark) e os schemas "
        "Pydantic",
        "screenshots/06_swagger.png", "http://localhost:8000/docs")
    screenshot_real(
        doc, "Endpoint de saúde da API — confirma API e banco operacionais",
        "screenshots/07_health.png", "http://localhost:8000/api/health")
    screenshot_real(
        doc, "Endpoint de comparativo de benchmark em 21/08/2026 — versão anterior à "
        "correção: todas as execuções divididas pelo sequencial mais recente (ver "
        "Seção 8.3)",
        "screenshots/08_comparativo.png",
        "http://localhost:8000/api/benchmark/comparativo")

    h2(doc, "9.4 Dashboard (frontend)")
    screenshot_real(
        doc, "Visão geral — badge \"API online\", KPIs de cobertura dos "
        "microdados (24.967 registros, 6 edições, 2 regiões, 16 UFs) e do "
        "benchmark, com o resumo por edição carregado da API",
        "screenshots/09_home.png", "http://localhost:5173")
    screenshot_real(
        doc, "Análises ENADE — evolução da Nota Geral por edição, com a barra "
        "de oito filtros combináveis e os indicadores do recorte",
        "screenshots/10_enade_anual.png", "http://localhost:5173/enade/anual")
    screenshot_real(
        doc, "Comparação longitudinal entre dois grupos (Norte × Nordeste) — "
        "gráfico, tabela ano a ano com o delta entre grupos e botões de "
        "exportação (CSV, SVG e relatório em Markdown)",
        "screenshots/11_comparar.png", "http://localhost:5173/enade/comparar")
    screenshot_real(
        doc, "Benchmark SPD em 21/08/2026 — os quatro gráficos e a interpretação gerada "
        "dos dados; o card '2,6157×' desta captura é o efeito do baseline errado "
        "corrigido na Seção 8.3",
        "screenshots/12_spd_comparativo.png",
        "http://localhost:5173/spd/comparativo")
    screenshot_real(
        doc, "Benchmark SPD — detalhamento da execução paralela com 2 workers: "
        "tempo por ano com cores por processo e tabela com worker_pid e "
        "timestamps de cada etapa",
        "screenshots/13_etapas.png", "http://localhost:5173/spd/etapas/2")

    h2(doc, "9.5 Build de produção do frontend")
    screenshot_real(
        doc, "npm run build — verificação de tipos (tsc) e build do Vite "
        "concluídos, com os artefatos gerados em dist/ (saída real do comando; "
        "o aviso de chunk é informativo, não é erro)",
        "screenshots/14_npm_build.png", "cd frontend; npm run build")


def secao_10_codigo(doc):
    h1(doc, "10. Principais trechos de código")

    par(doc,
        "Os oito quadros a seguir não cobrem operações básicas de CRUD — foram "
        "escolhidos por implementarem regras de negócio específicas, resolverem "
        "problemas concretos de engenharia ou materializarem, em código, as "
        "fórmulas centrais da fundamentação teórica de Sistemas Paralelos e "
        "Distribuídos.")

    code_quadro(
        doc,
        "Função pura e livre de efeitos colaterais para execução em processos "
        "paralelos",
        '''def processar_ano(ano: int) -> dict:
    """
    Executa o ETL de um ano e devolve métricas.

    SEM I/O em disco. SEM conexão de banco. SEM prints.
    Pode ser invocada por workers de ProcessPoolExecutor.

    Em caso de erro de leitura, devolve dict com status='erro' e mensagem
    em 'observacoes'. NÃO levanta exceção (para o pool continuar).
    """
    dt_ini, ts_ini = _agora()
    t0 = time.perf_counter()
    pid = os.getpid()

    out: dict = {
        "ano": ano, "worker_pid": pid,
        "timestamp_inicio": ts_ini, "timestamp_fim": ts_ini,
        "tempo_seg": 0.0, "status": "erro", "observacoes": "",
        # ... demais campos de métricas (linhas, médias, cursos únicos)
    }

    try:
        pasta = encontrar_pasta_ano(ano)
        if pasta is None:
            out["observacoes"] = f"Pasta 'microdados_enade_{ano}*' não existe"
            return _fechar(out, t0)
        # ... leitura de arq1/arq3, filtro do recorte, agregação de notas ...
        out["status"] = "ok"
    except Exception as exc:                          # nunca quebra o pool
        out["status"] = "erro"
        out["observacoes"] = f"{type(exc).__name__}: {exc}"

    return _fechar(out, t0)


def _fechar(out: dict, t0: float) -> dict:
    dt_fim, ts_fim = _agora()
    out["timestamp_fim"] = ts_fim
    out["tempo_seg"] = round(time.perf_counter() - t0, 4)
    return out''',
        "no Windows, multiprocessing usa o método spawn — cada processo novo "
        "reimporta o módulo do zero, em vez de herdar memória via fork. Por "
        "isso a função precisa ser importável no nível do módulo (top-level), "
        "sem conexão de banco (cada worker não pode competir pelo pool de "
        "conexões do processo principal) e nunca deixar uma exceção escapar "
        "(um ano com erro não pode derrubar os outros cinco que estão sendo "
        "processados em paralelo).",
        fonte="etl/processar_ano.py")

    code_quadro(
        doc,
        "Distribuição do trabalho entre processos com ProcessPoolExecutor",
        '''mem_inicio = memoria_atual_mb()
t0 = time.perf_counter()
etapas: list[dict] = []
mem_pico = mem_inicio

def ordem_de_submissao(ordem, tempos_base):
    if ordem == "crescente":
        return sorted(ANOS)                      # 2005 ... 2021 (ordem histórica)
    base = tempos_base or tempos_proxy_tamanho() # LPT: maior tempo primeiro (Graham)
    return sorted(ANOS, key=lambda a: (-base.get(a, 0.0), a))

ordem_anos = ordem_de_submissao(ordem, tempos_base)
# ProcessPoolExecutor entrega cada tarefa ao PRIMEIRO worker livre, na ordem
# de submissão — ele não balanceia. A ordem, portanto, define o makespan.
with ProcessPoolExecutor(max_workers=num_workers) as executor:
    futuros = {executor.submit(processar_ano, ano): ano for ano in ordem_anos}
    for fut in as_completed(futuros):
        ano = futuros[fut]
        try:
            et = fut.result()
        except Exception as exc:                      # nunca quebra a suíte
            et = {"ano": ano, "worker_pid": -1, "status": "erro",
                  "observacoes": f"future: {exc}"}      # ... demais campos
        etapas.append(et)
        mem_pico = max(mem_pico, memoria_atual_mb())

etapas.sort(key=lambda e: e["ano"])''',
        "implementa o padrão fork-join do experimento e torna a ORDEM de "
        "submissão uma variável controlada: o executor entrega cada ano ao "
        "primeiro worker livre, então com 6 unidades de tamanhos diferentes a "
        "ordem muda o makespan sem tocar no código dos workers. A ordem "
        "'crescente' reproduz o comportamento histórico; 'lpt' submete os anos "
        "em ordem decrescente de tempo (Graham, 1969), usando os tempos por ano "
        "do sequencial da mesma suíte. A coleta continua por as_completed, com "
        "reordenação determinística antes de gravar.",
        fonte="scripts/09_benchmark_paralelo.py")

    code_quadro(
        doc,
        "Cálculo de speedup e eficiência como view SQL auditável",
        '''CREATE OR REPLACE VIEW v_benchmark_metricas AS
SELECT
    p.id                 AS execucao_id,
    p.num_workers,
    s.tempo_total_seg    AS tempo_sequencial,
    p.tempo_total_seg    AS tempo_paralelo,
    ROUND((s.tempo_total_seg / p.tempo_total_seg)::NUMERIC, 4) AS speedup,
    ROUND(((s.tempo_total_seg / p.tempo_total_seg)
           / p.num_workers)::NUMERIC, 4)                     AS eficiencia,
    p.suite_id, p.campanha_id, p.ordem_submissao,
    s.id                 AS baseline_execucao_id,
    CASE WHEN p.suite_id IS NULL THEN 'temporal' ELSE 'suite' END AS pareamento
FROM benchmark_execucao p
LEFT JOIN LATERAL (
    SELECT b.* FROM benchmark_execucao b
    WHERE b.modo = 'sequencial' AND b.aquecimento = FALSE
      AND ((p.suite_id IS NOT NULL AND b.suite_id = p.suite_id)
        OR (p.suite_id IS NULL AND b.timestamp_inicio < p.timestamp_inicio))
    ORDER BY b.timestamp_inicio DESC LIMIT 1
) s ON TRUE
WHERE p.modo = 'paralelo';''',
        "materializa em SQL as duas fórmulas centrais — S(p) = T(1)/T(p) e "
        "E(p) = S(p)/p — pareando cada execução paralela com o sequencial da "
        "PRÓPRIA suíte (mesmo suite_id). Linhas antigas, gravadas antes de "
        "existir o conceito de suíte, mantêm o pareamento temporal com o "
        "sequencial imediatamente anterior, e a coluna pareamento diz qual "
        "regra foi usada. É a única definição de speedup do sistema: API, "
        "dashboard, documento e slides leem esta view.",
        fonte="scripts/14_migrar_schema_v2.py")

    code_quadro(
        doc,
        "Checagem cruzada: a view recalculada em Python a partir das tabelas",
        '''def parear_em_python(execs):
    # Reproduz a regra da view: mesma suíte; senão, o sequencial anterior.
    seqs = [e for e in execs if e["modo"] == "sequencial" and not e["aquecimento"]]
    out = {}
    for p in execs:
        if p["modo"] != "paralelo":
            continue
        if p["suite_id"] is not None:
            cands = [s for s in seqs if s["suite_id"] == p["suite_id"]]
        else:
            cands = [s for s in seqs if s["timestamp_inicio"] < p["timestamp_inicio"]]
        out[p["id"]] = max(cands, key=lambda s: s["timestamp_inicio"]) if cands else None
    return out

# para cada paralela: speedup_view == round(t_seq_pareado / t_par, 4)
s_py = r4(Decimal(str(base["tempo_total_seg"])) / Decimal(str(p["tempo_total_seg"])))
check(Decimal(str(v["speedup"])) == s_py and v["baseline_execucao_id"] == base["id"], ...)''',
        "a versão anterior do endpoint /api/benchmark/comparativo recalculava "
        "speedup em Python contra 'o sequencial mais recente' — e, quando o "
        "banco passou a conter duas rodadas, dividiu uma execução de 21/06 "
        "(43,17 s) pelo baseline de 21/08 (112,93 s), exibindo 2,6157× onde o "
        "valor pareado é 2,0749×. A correção inverteu os papéis: a API só lê a "
        "view; o Python recalcula tudo a partir das tabelas apenas para "
        "CONFERIR (scripts/13_validar_metricas.py e tests/), e falha se um "
        "único número divergir.",
        fonte="scripts/13_validar_metricas.py")

    code_quadro(
        doc,
        "Composição segura de filtros SQL dinâmicos",
        '''@dataclass
class FiltrosAnalise:
    nu_ano: int | None = None
    co_regiao: int | None = None
    co_uf: int | None = None
    co_ies: int | None = None
    # ... demais filtros (grupo, modalidade, categ. adm., org. acad.)

    def to_where(self, alias: str | None = None) -> tuple[str, dict]:
        prefix = f"{alias}." if alias else ""
        fragments: list[str] = []
        params: dict = {}
        for key, col in _FILTRO_COLUNAS.items():
            value = getattr(self, key)
            if value is not None:
                fragments.append(f"AND {prefix}{col} = %({key})s")
                params[key] = value
        return (" ".join(fragments), params)''',
        "regra reutilizada por todos os endpoints de análise (resumo por "
        "ano, região, UF, IES e registros paginados): monta a cláusula "
        "WHERE dinamicamente a partir de até oito filtros combináveis, "
        "sempre com placeholders nomeados (%(chave)s) — nunca "
        "concatenação direta de valor na string SQL —, eliminando a "
        "classe inteira de vulnerabilidades de injeção de SQL nesse ponto "
        "da API.",
        fonte="api/dependencies.py")

    code_quadro(
        doc,
        "Correção de um bug de formatação na carga em lote",
        '''def preparar_buffer(df: pd.DataFrame) -> io.StringIO:
    out = df[COLUNAS_ORIGEM].copy()

    # ---- correção do bug "2806701.000000" ----
    # to_numeric tolera valores já int/float; astype('Int64') aceita NaN
    # e produz <NA>, que to_csv com na_rep='' transforma em string vazia.
    for col in COLUNAS_INTEIRAS_ORIGEM:
        out[col] = pd.to_numeric(out[col], errors="coerce").astype("Int64")

    for col in COLUNAS_NOTAS_ORIGEM:
        out[col] = pd.to_numeric(out[col], errors="coerce")

    buf = io.StringIO()
    out.to_csv(buf, sep=";", header=False, index=False,
               na_rep="", float_format="%.4f")
    buf.seek(0)
    return buf''',
        "colunas inteiras com valores nulos (ex.: código do município) "
        "chegam ao pandas como float64, e sem tratamento seriam gravadas "
        "como \"2806701.000000\" em vez de \"2806701\" — incompatível com "
        "as colunas INTEGER/SMALLINT do schema. A conversão para o tipo "
        "nullable Int64 resolve nulo e formatação ao mesmo tempo, mantendo "
        "o COPY FROM STDIN como caminho de carga (ordens de grandeza mais "
        "rápido que INSERT linha a linha).",
        fonte="scripts/06_carregar_dados_postgres.py")

    code_quadro(
        doc,
        "Interpretação do benchmark gerada a partir dos dados, não fixa no "
        "código",
        '''// Métrica de Karp–Flatt: estima a fração sequencial do programa a
// partir do speedup medido S com p processos —
// e = (1/S − 1/p) / (1 − 1/p).
function fracaoSerialKarpFlatt(speedup: number, p: number): number | null {
  if (p <= 1 || speedup <= 0) return null;
  const e = (1 / speedup - 1 / p) / (1 - 1 / p);
  return Number.isFinite(e) ? e : null;
}

// A narrativa muda conforme os dados realmente carregados no banco —
// nunca afirma "o paralelismo saturou" se o maior número de processos
// testado ainda foi o mais rápido.
const ordenados = [...paralelos].sort((a, b) => a.num_workers - b.num_workers);
const maxWorkers = ordenados[ordenados.length - 1].num_workers;
const escalouAteOFim = melhor.num_workers === maxWorkers;''',
        "a tela de comparativo do benchmark inicialmente trazia um texto "
        "fixo (\"2 processos superou 4\"), redigido para a rodada oficial "
        "do experimento. Ao recarregar outra rodada — na qual 4 processos "
        "venceu —, o texto passou a contradizer os próprios gráficos da "
        "mesma tela. A correção substitui o texto fixo por uma "
        "interpretação derivada dos dados carregados, incluindo a "
        "estimativa de Karp–Flatt da fração sequencial por configuração — "
        "com a ressalva, explicitada na Seção 8.3, de que essa estimativa "
        "embute granularidade, escalonamento, contenção e overhead num único "
        "número.",
        fonte="frontend/src/pages/spd/Comparativo.tsx")

    code_quadro(
        doc,
        "Autodetecção de rota de conexão com um serviço externo",
        '''def candidatos_conexao(senha: str, dbname: str, port: int
                       ) -> list[tuple[str, dict]]:
    """Rotas a tentar, em ordem, quando o host não é fixado no ambiente.

    O prefixo do cluster da pooler (aws-0/aws-1/...) varia por projeto e
    não é derivável da região — chutar errado dá `FATAL (ENOTFOUND)
    tenant/user ... not found`. Tentamos os clusters conhecidos e, por
    último, o host direto (IPv6-only).
    """
    base = {"dbname": dbname, "password": senha, "sslmode": "require",
            "connect_timeout": 20}
    usuario_pooler = f"postgres.{PROJECT_REF}"
    rotas = []
    for cluster in ("aws-1", "aws-0"):
        host = f"{cluster}-{REGIAO}.pooler.supabase.com"
        rotas.append((
            f"postgresql://{usuario_pooler}@{host}:{port}/{dbname}",
            {**base, "host": host, "port": port, "user": usuario_pooler},
        ))
    host_direto = f"db.{PROJECT_REF}.supabase.co"
    rotas.append((
        f"postgresql://postgres@{host_direto}:5432/{dbname} (IPv6)",
        {**base, "host": host_direto, "port": 5432, "user": "postgres"},
    ))
    return rotas''',
        "ao integrar a base consolidada a um projeto Supabase (réplica em "
        "nuvem para o ecossistema EnadeX), a primeira tentativa de conexão "
        "falhou com um erro de \"tenant not found\" porque o prefixo do "
        "cluster da pooler não é previsível a partir da região do projeto. "
        "Em vez de fixar um valor por tentativa e erro, a função tenta as "
        "rotas conhecidas em ordem, com timeout curto por tentativa — o "
        "mesmo princípio de tolerância a falhas do restante do pipeline.",
        fonte="scripts/11_carregar_supabase.py")


def secao_fontes(doc):
    h1(doc, "Fontes e documentação de apoio")
    par(doc,
        "Este documento foi redigido a partir da documentação técnica mantida "
        "no próprio repositório do projeto e do código-fonte efetivamente "
        "implementado — nenhum número ou trecho de código apresentado nas "
        "seções anteriores é hipotético. As principais fontes internas "
        "consultadas foram:")
    bullets(doc, [
        "README.md — visão geral, stack e roteiro de execução do projeto.",
        "docs/ARQUITETURA.md — camadas, fluxo de dados e decisões técnicas.",
        "docs/DESIGN_LOG.md — histórico cronológico de decisões (D1–D17).",
        "docs/RELATORIO_VERIFICACAO.md — estado real do projeto, do EnadeX e do Supabase em 2026-09-01.",
        "Banco de dados local: views v_benchmark_metricas e v_benchmark_resumo (fonte de todos os números da Seção 8.3).",
        "docs/DICIONARIO_VARIAVEIS.md — harmonização variável a variável entre edições.",
        "docs/LOG_EXECUCAO_ETL.md — versão do dataset e contagens por ano.",
        "docs/RESULTADOS_BENCHMARK.md — números oficiais do experimento de paralelismo.",
        "docs/INTEGRACAO_ECOSSISTEMA.md — contrato de consumo da base pelo ecossistema EnadeX.",
        "docs/GUIA_APRESENTACAO.md — roteiro oral e perguntas frequentes da banca.",
        "docs/EVIDENCIAS_TESTES.md — checklist de evidências, base da Seção 9.",
        "docs/PROXIMOS_PASSOS.md — roadmap e itens conscientemente fora do escopo atual.",
        "Código-fonte de api/, frontend/src/, scripts/ e etl/.",
    ])


# ===========================================================================
# MONTAGEM DO DOCUMENTO
# ===========================================================================

def configurar_estilo_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = CINZA_TEXTO
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    for nivel, tamanho in ((1, 18), (2, 14), (3, 12)):
        style = doc.styles[f"Heading {nivel}"]
        style.font.name = "Calibri"
        style.font.size = Pt(tamanho)
        style.font.bold = True


def main():
    doc = Document()
    configurar_estilo_base(doc)
    core = doc.core_properties
    core.title = "ENADE-Time Distribuído — Documento Técnico-Acadêmico"
    core.author = "Lucas Eduardo Tavares Costa"
    core.subject = "Sistema Paralelo de Análise Longitudinal dos Microdados do ENADE"

    secao_capa(doc)
    secao_resumo(doc)
    secao_1_problema(doc)
    secao_2_objetivo_geral(doc)
    secao_3_objetivos_especificos(doc)
    secao_4_justificativa(doc)
    secao_5_ferramentas(doc)
    secao_6_processo(doc)
    secao_7_arquitetura_geral(doc)
    secao_8_integracao_subprojeto(doc)
    secao_8_3_resultados(doc)
    secao_9_capturas(doc)
    secao_10_codigo(doc)
    secao_fontes(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUT_DIR / "ENADE_Time_Distribuido_Documento_Academico.docx"
    try:
        doc.save(str(out_path))
    except PermissionError:
        # arquivo aberto no Word — salva ao lado com sufixo para não perder o trabalho
        from datetime import datetime
        alt = OUT_DIR / (out_path.stem + datetime.now().strftime("_%H%M") + ".docx")
        doc.save(str(alt))
        print(f"AVISO: '{out_path.name}' está aberto no Word (lock). "
              f"Salvo como: {alt.name}")
        out_path = alt
    print(f"Documento salvo em: {out_path}")
    print(f"Tamanho: {out_path.stat().st_size / 1024:.1f} KB")
    print(f"Quadros de código: {_QUADRO_N[0]}")
    print(f"Figuras: {_FIGURA_N[0]}")
    print(f"Tabelas: {_TABELA_N[0]}")


if __name__ == "__main__":
    main()
